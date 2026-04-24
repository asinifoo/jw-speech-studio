"""md 파싱, 텍스트 파싱, DOCX 파싱, 유형 코드, 검증"""
import io
import os
import re
from config import _OUTLINES_DIR
from services.bible_utils import extract_scriptures_from_text


def _split_comma_refs(text):
    """쉼표로 참조 분리하되, 따옴표("" 또는 "") 안의 쉼표는 무시."""
    parts = []
    buf = []
    in_q = False
    for i, ch in enumerate(text):
        if ch in ('"', '\u201c'):
            in_q = True
        elif ch in ('"', '\u201d'):
            in_q = False
        elif ch == ',' and not in_q:
            rest = text[i + 1:].lstrip()
            if rest and ('\uac00' <= rest[0] <= '\ud7a3' or rest[0] == '\u300c'):
                parts.append(''.join(buf).strip())
                buf = []
                continue
        buf.append(ch)
    tail = ''.join(buf).strip()
    if tail:
        parts.append(tail)
    return parts

_TYPE_NAMES = {
    "S-34": "공개강연", "S-31": "기념식", "S-123": "특별강연", "S-211": "RP모임",
    "SB": "생활과봉사",
    "CO": "대회", "CO_C": "대회(순회)", "CO_R": "대회(지역)",
    "JWBC": "JW방송", "JWBC-SP": "JW방송(연설)", "JWBC-MW": "JW방송(아침숭배)", "JWBC-PG": "JW방송(월간프로그램)", "JWBC-AM": "JW방송(연례총회)",
    "ETC": "기타",
}

_OUTLINE_TYPE_KO_TO_EN = {
    "공개강연": "S-34", "공개 강연": "S-34",
    "기념식": "S-31",
    "특별강연": "S-123", "특별 강연": "S-123",
    "RP모임": "S-211", "RP 모임": "S-211",
    "순회대회": "CO_C", "순회 대회": "CO_C",
    "지역대회": "CO_R", "지역 대회": "CO_R",
    "생활과봉사": "SB", "생활과 봉사": "SB",
    "기타": "ETC",
}

# /api/outline/types 응답 메타 (Phase 1 Step 2a).
# _TYPE_NAMES 는 name 단일 소스 (영→한글). _TYPE_META 는 나머지 메타 (aliases/패턴/예시).
# _TYPE_NAMES "CO" 는 Gather.jsx OUTLINE_TYPES 의 wrapper(code='CO', sub=[CO_C,CO_R]) 용
# fallback 으로 유지되나, 저장용 코드는 아니므로 get_outline_types() 응답에서 제외.
_TYPE_META = {
    "S-34": {
        "aliases": ["공개 강연"],
        "num_pattern": "001~194+ (시리즈 번호)",
        "version_example": "10/24 (발행 월/년)",
        "year_required": False,
    },
    "S-31": {
        "aliases": [],
        "num_pattern": "001 시작 (필요 시 002, 003)",
        "version_example": "8/19 (개정 월/년)",
        "year_required": False,
    },
    "S-123": {
        "aliases": ["특별 강연"],
        "num_pattern": "001 시작 (필요 시 002, 003)",
        "version_example": "5/26 (발표 월/년)",
        "year_required": False,
    },
    "S-211": {
        "aliases": ["RP 모임"],
        "num_pattern": "001 시작 (필요 시 002, 003)",
        "version_example": "6/26 (개최 월/년)",
        "year_required": False,
    },
    "SB": {
        "aliases": ["생활과 봉사"],
        "num_pattern": "MMW (월+주차, 예: 041 = 4월 1주차)",
        "version_example": "4/26 (사용 월/년)",
        "year_required": False,
    },
    "CO_C": {
        "aliases": ["순회대회", "순회 대회"],
        "num_pattern": "001 시작 (상/하반기 또는 추가)",
        "version_example": "3/26 (개최 월/년)",
        "year_required": False,
    },
    "CO_R": {
        "aliases": ["지역대회", "지역 대회"],
        "num_pattern": "001 시작 (필요 시 002, 003)",
        "version_example": "7/26 (개최 월/년)",
        "year_required": False,
    },
    "JWBC": {
        "aliases": [],
        "num_pattern": "(5d 설계 예정)",
        "version_example": "",
        "year_required": False,
    },
    "JWBC-SP": {
        "aliases": [],
        "num_pattern": "(5d 설계 예정)",
        "version_example": "",
        "year_required": False,
    },
    "JWBC-MW": {
        "aliases": [],
        "num_pattern": "(5d 설계 예정)",
        "version_example": "",
        "year_required": False,
    },
    "JWBC-PG": {
        "aliases": [],
        "num_pattern": "(5d 설계 예정)",
        "version_example": "",
        "year_required": False,
    },
    "JWBC-AM": {
        "aliases": [],
        "num_pattern": "(5d 설계 예정)",
        "version_example": "",
        "year_required": False,
    },
    "ETC": {
        "aliases": [],
        "num_pattern": "",
        "version_example": "",
        "year_required": False,
    },
}

# `_TYPE_NAMES` 에는 있지만 `_TYPE_META` 에는 의도적으로 제외된 코드.
# Gather.jsx OUTLINE_TYPES 의 wrapper 용 fallback (저장 경로엔 안 씀).
_TYPE_META_EXCLUDED = {"CO"}


def get_outline_types() -> list:
    """GET /api/outline/types 응답 데이터.

    _TYPE_NAMES (name 단일 소스) + _TYPE_META (aliases/패턴/예시) 병합.
    _TYPE_META_EXCLUDED 의 코드는 응답에서 제외 (Gather wrapper 용).
    """
    result = []
    for code, name in _TYPE_NAMES.items():
        if code in _TYPE_META_EXCLUDED:
            continue
        meta = _TYPE_META.get(code, {})
        result.append({
            "code": code,
            "name": name,
            "aliases": list(meta.get("aliases", [])),
            "num_pattern": meta.get("num_pattern", ""),
            "version_example": meta.get("version_example", ""),
            "year_required": bool(meta.get("year_required", False)),
        })
    return result


def normalize_outline_type(value: str) -> str:
    """한글/영문 혼재 outline_type을 영문 코드로 정규화.

    - 매핑 테이블 일치 → 영문 반환
    - 영문 prefix (S-*/CO_*/SB/JWBC*/ETC) → 그대로
    - 빈 값 → 빈 문자열
    - 알 수 없음 → 원본 반환 (호출측에서 판단)
    """
    if not value:
        return ""
    v = value.strip()
    if v in _OUTLINE_TYPE_KO_TO_EN:
        return _OUTLINE_TYPE_KO_TO_EN[v]
    if v.startswith("S-") or v.startswith("CO_") or v == "SB" or v == "ETC" or v.startswith("JWBC"):
        return v
    return v


def _outline_prefix(otype: str, onum: str, year: str = "") -> str:
    """유형코드+번호+year → 파일/ID용 prefix

    year 있으면 '{code}_{num}_y{year}' 형태, 없으면 '{code}_{num}'.
    year="" 기본값으로 기존 호출처 무수정 호환.
    ETC/빈 type은 year 무시 (기존 동작 유지).
    """
    # 숫자 번호면 3자리 패딩
    num = onum.zfill(3) if onum.isdigit() else onum
    year_tag = f"_y{year}" if year else ""

    if otype in ("공개강연",) or otype.startswith("S-34"):
        return f"S-34_{num}{year_tag}"
    elif otype in ("기념식",) or otype.startswith("S-31"):
        return f"S-31_{num}{year_tag}"
    elif otype.startswith("JWBC"):
        return f"{otype}_{num}{year_tag}"
    elif otype.startswith("S-") or otype.startswith("CO") or otype.startswith("SB"):
        return f"{otype}_{num}{year_tag}"
    elif otype == "ETC" or not otype:
        return onum
    else:
        return f"{otype}_{num}{year_tag}"


def _ver_safe(version: str) -> str:
    """버전 문자열을 파일명에 안전하게"""
    return version.replace("/", "-").replace(" ", "").strip()


_TRAILING_MARKER_RE = re.compile(
    r'(\s*(?:\[\s*시각\s*자료\s*\d+\s*\]|\[\s*지시문\s*\]|\[\s*연사\s*지시\s*\]|\[\s*영상\s*\d+\s*\]|\[\s*낭독\s*\])\s*)+$'
)
# 대괄호 낭독 성구: '[이사야 46:9, 10 낭독]', '[시편 1:2, 3 및 각주 낭독]',
# '[요한 1서 2:4 낭독]' (신세계역), '[요한 첫째 2:4 낭독]' (구판), '[베드로 후서 3:13 낭독]' (공백 포함)
# group(1) = 성구 부분만 (낭독 제외). 책명 부분은 '한글 단어' 또는 'N서' 토큰이 공백으로 이어지는 형태.
_BRACKET_READING_RE = re.compile(
    r'\[\s*([가-힣]+(?:\s+(?:[가-힣]+|\d+서))*\s+\d+:[\d,\s\-]+(?:\s*(?:및\s*)?각주)?)\s*낭독\s*\]'
)


def parse_outline_text(text: str, has_separate_title: bool = False) -> dict:
    lines = [l.strip() for l in text.strip().split("\n") if l.strip()]
    if not lines:
        return {"title": "", "duration": "", "points": []}

    if has_separate_title:
        title = ""
        duration = ""
        point_lines = lines
    else:
        first_line = lines[0]
        duration = ""
        title = ""

        dur_match = re.search(r"\((\d+)\s*분\)", first_line)
        has_refs = re.search(r"\([^)]*(?:사|마|눅|요|창|출|레|민|신|삿|룻|삼|왕|대|라|느|더|욥|시|잠|전|아|렘|애|겔|단|호|욜|암|옵|욘|미|나|합|습|학|슥|말|막|행|롬|고|갈|엡|빌|골|살|딤|딛|몬|히|약|벧|유|계|「)\s", first_line)

        if dur_match and not has_refs:
            duration = dur_match.group(1) + "분"
            title = re.sub(r"\(\d+\s*분\)", "", first_line).strip()
            point_lines = lines[1:]
        elif not has_refs and len(lines) > 1:
            title = first_line
            point_lines = lines[1:]
        else:
            title = ""
            point_lines = lines

    points = []
    for raw_line in point_lines:
        line = raw_line.strip()
        if not line:
            continue

        dash_match = re.match(r'^(-{1,4})(.+)$', line)
        if dash_match:
            level = len(dash_match.group(1)) + 1
            line = dash_match.group(2).strip()
        elif raw_line.startswith('  ') or raw_line.startswith('\t'):
            level = 2
        else:
            level = 1

        # 1) 대괄호 낭독 성구 추출 + [낭독]으로 치환. 괄호 성구와 일관되게 scriptures로 이동.
        #    치환 후 [낭독]은 trailing 마커로 취급되어 2)에서 분리됨.
        bracket_scriptures = []
        def _br_sub(m):
            bracket_scriptures.append(f"{m.group(1).strip()} (낭독)")
            return "[낭독]"
        line = _BRACKET_READING_RE.sub(_br_sub, line)

        # 2) 줄 끝 마커 분리 ([시각 자료 N] / [지시문] / [연사 지시] / [영상 N] / [낭독])
        #    기존 괄호 성구 추출이 줄 끝 기준이라 마커가 붙어있으면 매치 실패 → 미리 떼어냄
        trailing_markers = ''
        mm = _TRAILING_MARKER_RE.search(line)
        if mm:
            trailing_markers = ' ' + line[mm.start():].strip()
            line = line[:mm.start()].rstrip()

        refs_match = re.search(r"\(([^)]+)\)\s*$", line)
        scriptures = list(bracket_scriptures)  # 낭독 성구 먼저
        publications = []
        point_text = line

        if refs_match:
            refs_str = refs_match.group(1)
            point_text = line[:refs_match.start()].strip()
            parts = re.split(r";\s*", refs_str)
            expanded_parts = []
            for part in parts:
                expanded_parts.extend(_split_comma_refs(part))
            for part in expanded_parts:
                part = part.strip()
                if part.startswith("\u300c") or part.startswith("'"):
                    publications.append(part)
                else:
                    scriptures.append(part)

        body_refs = extract_scriptures_from_text(point_text)
        for br in body_refs:
            if br not in scriptures:
                scriptures.append(br)

        # 마커 복원
        if trailing_markers:
            point_text = point_text + trailing_markers

        point_data = {"title": point_text, "scriptures": scriptures, "publications": publications, "level": level}

        if level >= 2 and points:
            if "sub_points" not in points[-1]:
                points[-1]["sub_points"] = []
            points[-1]["sub_points"].append(point_data)
        else:
            points.append(point_data)

    if not points:
        body_refs = extract_scriptures_from_text(title)
        points.append({"title": title, "scriptures": body_refs, "publications": []})

    return {"title": title, "duration": duration, "points": points}


# ─── DOCX 파서 (결정론적) ──────────────────────────────────

def _docx_normalize_text(text: str) -> str:
    text = text.replace('\n', ' ').replace('\r', ' ')
    text = text.replace('\u200b', '').replace('\u2002', ' ').replace('\xa0', ' ')
    text = re.sub(r'[\uf000-\uf8ff]', '', text)
    return text.strip()


def _docx_get_level(indent_pt: float) -> int:
    if indent_pt <= 10: return 1
    elif indent_pt <= 25: return 2
    elif indent_pt <= 40: return 3
    elif indent_pt <= 55: return 4
    elif indent_pt <= 70: return 5
    else: return 6


def _docx_first_run_bold(para) -> bool:
    for run in para.runs:
        if run.text.strip():
            return run.bold == True
    return False


def _docx_any_italic(para) -> bool:
    for run in para.runs:
        if run.text.strip() and run.italic:
            return True
    return False


def _docx_classify_line(text: str, indent_pt: float, bold: bool, italic: bool, style_name: str = "") -> str:
    """줄을 분류해 태그 문자열 반환. L1~L5 | 소주제 | 제목 | 주의 | 하단 | 제외"""
    # 노래/기도 → 제외
    if re.match(r'^노래\s+\d+', text):
        return "제외"
    # 하단 문서번호
    if re.match(r'^S-\d+', text):
        return "하단"
    # 총 시간 안내 ("30분에 다룰 것" / "30분을 다룰 것")
    if re.search(r'\d+\s*분\s*(?:을|에)?\s*다룰\s*것', text):
        return "하단"
    # 하단 안내문
    if '골자에 밀접히 고착' in text:
        return "하단"
    # 저작권
    if text.startswith('©') or 'Watch Tower' in text:
        return "하단"
    # 스타일 기반 하단
    if style_name in ('JW TOTAL TIME', 'JW Copyright', 'JW Mnemonic'):
        return "하단"
    # 스타일 기반 제목 분류
    if style_name.startswith('JW Title'):
        return "제목"
    # 소주제
    if style_name.startswith('JW HEADING'):
        return "소주제"
    if indent_pt <= 10 and re.search(r'\(\s*\d+\s*분\s*\)', text):
        return "소주제"
    # 유의 사항
    if '유의 사항' in text or style_name == 'JW Note':
        if '골자에 밀접히 고착' not in text:
            return "주의"
    if italic and indent_pt > 10 and indent_pt <= 15:
        return "주의"
    # 제목 (볼드, 들여쓰기 0~10pt, 시간 없음)
    if bold and indent_pt <= 10 and not re.search(r'\(\s*\d+\s*분\s*\)', text):
        return "제목"
    # 일반 요점
    return f"L{_docx_get_level(indent_pt)}"


def _docx_merge_split_notes(lines: list) -> list:
    """'연사의 유의 사항:' 뒤 실제 내용이 다음 줄로 분리된 경우 합치기"""
    merged = []
    i = 0
    while i < len(lines):
        ln = lines[i]
        if ln["tag"] == "주의":
            clean = re.sub(r'^(연사의\s*)?유의\s*사항\s*:?\s*', '', ln["text"]).strip()
            if not clean and i + 1 < len(lines):
                nxt = lines[i + 1]
                if nxt["tag"] in ("L1", "L2"):
                    combined = dict(ln)
                    combined["text"] = ln["text"].rstrip() + " " + nxt["text"].strip()
                    merged.append(combined)
                    i += 2
                    continue
        merged.append(ln)
        i += 1
    return merged


_PUB_TAIL_RE = re.compile(r'^[\d\-,\s]+(면|쪽|항)(\s+[\d\-,\s]+(면|쪽|항))?\s*\)?$')
# 독립 성구 줄 (완전한 괄호): '(고전 13:4-7)', '(요1 5:3)', '(시 4:22-24; 고전 15:33)'
# 보호: '(30분에 다룰 것)' 같이 '분'만 있는 패턴, '(「파10」 ...)' 출판물 괄호는 매치 X
_STANDALONE_SCRIPTURE_RE = re.compile(r'^\(\s*[가-힣][가-힣\s\d]*\s+\d+:\d+.*\)$')


def _docx_merge_scripture_linebreaks(lines: list) -> list:
    """DOCX 줄바꿈으로 쪼개진 성구/출판물/독립 괄호 병합.

    병합 우선순위:
    1. 성구 꼬리 (예: '(벧후' + '3:13)', '10:1-3)', '3:11, 12)')
    2. 출판물 꼬리 (예: '(「깨」 15/6' + '8-9면)', '5면 3항)', '22-24면 10-11항)')
    3. 독립 성구 괄호 (예: '...것이다' + '(고전 13:4-7)')
       — 레벨 보호: 이전 줄이 L 태그이고 현재 줄 레벨이 이전 이상일 때만 병합
    """
    merged = []
    for ln in lines:
        text_only = ln["text"].strip()
        is_scripture_tail = (
            re.match(r'^[\d:,;\s\-]+\)?', text_only)
            and not re.search(r'[가-힣a-zA-Z]', text_only)
        )
        is_publication_tail = _PUB_TAIL_RE.match(text_only) is not None
        is_standalone_scripture = _STANDALONE_SCRIPTURE_RE.match(text_only) is not None

        should_merge = False
        if merged:
            if is_scripture_tail or is_publication_tail:
                should_merge = True
            elif is_standalone_scripture:
                prev_tag = merged[-1].get("tag", "")
                cur_tag = ln.get("tag", "")
                if prev_tag.startswith("L") and cur_tag.startswith("L"):
                    try:
                        if int(cur_tag[1:]) >= int(prev_tag[1:]):
                            should_merge = True
                    except ValueError:
                        pass

        if should_merge:
            prev = dict(merged[-1])
            prev["text"] = prev["text"] + " " + text_only
            merged[-1] = prev
        else:
            merged.append(ln)
    return merged


_REF_PAREN_RE = re.compile(r'\(([^)]+)\)\s*$')
_TIME_ONLY_RE = re.compile(r'^\s*\d+\s*분\s*$')
_READING_RE = re.compile(r'\[\s*낭독\s+([^\]]+)\]')


def _docx_extract_refs(text: str):
    """텍스트에서 성구/출판물 추출. content, scriptures, publications 반환."""
    scriptures = []
    publications = []
    content = text

    # [낭독 사 48:17] 패턴 — 성구로 추출하되 content에는 보존
    for m in _READING_RE.finditer(content):
        ref = m.group(1).strip()
        if ref and ref not in scriptures:
            scriptures.append(ref)

    # 끝 괄호 (...) 안 성구/출판물 분리
    m = _REF_PAREN_RE.search(content)
    if m:
        refs_str = m.group(1).strip()
        if not _TIME_ONLY_RE.match(refs_str):
            content_before = content[:m.start()].rstrip()
            parts = re.split(r'[;；]', refs_str)
            kept = True
            for part in parts:
                part = part.strip()
                if not part:
                    continue
                if '「' in part or '」' in part:
                    publications.append(part)
                elif re.search(r'[가-힣a-zA-Z]', part) and re.search(r'\d', part):
                    # 성구 후보: 한글/영문 + 숫자
                    if part not in scriptures:
                        scriptures.append(part)
                else:
                    # 성구도 출판물도 아니면 content에 유지
                    kept = False
                    break
            if kept and (scriptures or publications):
                content = content_before

    # 본문의 성구 추출 (중복 제거)
    body_refs = extract_scriptures_from_text(content)
    for br in body_refs:
        if br not in scriptures:
            scriptures.append(br)

    return content, scriptures, publications


def _docx_assign_numbers(points: list) -> None:
    """points 리스트에 parent 기반 계층 번호 부여 (1, 1.1, 1.1.1, 1.2, 2, ...)"""
    path = []            # 현재 경로의 sibling 번호 (L1, L2, ...)
    sibling = {}         # parent tuple → 해당 부모 밑의 sibling 개수

    for pt in points:
        level = pt["level"]
        if level < 1:
            level = 1
            pt["level"] = 1

        while len(path) >= level:
            path.pop()

        parent = tuple(path)
        sibling[parent] = sibling.get(parent, 0) + 1
        path.append(sibling[parent])
        pt["number"] = ".".join(str(n) for n in path)

        # 이 위치 아래에 남아있는 오래된 카운터는 제거 (새 분기 시작)
        to_del = [k for k in sibling if len(k) > len(path) and k[:len(path)] == tuple(path)]
        for k in to_del:
            del sibling[k]


def parse_outline_docx(file_bytes: bytes) -> dict:
    """
    골자 DOCX bytes를 받아서 구조화된 dict를 반환한다.
    결정론적 파싱 — LLM을 사용하지 않는다.
    """
    try:
        from docx import Document
    except ImportError as e:
        raise RuntimeError("python-docx가 설치되지 않았습니다. pip install python-docx") from e

    try:
        doc = Document(io.BytesIO(file_bytes))
    except Exception as e:
        raise ValueError(f"DOCX 파일을 열 수 없습니다: {e}")

    # 구버전 포맷 검증
    styles = set()
    for para in doc.paragraphs:
        if para.text.strip() and para.style:
            styles.add(para.style.name)
    supported = any(s.startswith('JW') or s == 'Normal' for s in styles)
    if not supported:
        short_styles = [s for s in styles if len(s) <= 2]
        if styles and len(short_styles) == len(styles):
            raise ValueError(
                f"지원하지 않는 구버전 DOCX 포맷입니다. 스타일: {sorted(styles)}. "
                f"성구 콜론(:)이 손실된 파일은 처리할 수 없습니다. "
                f"신 포맷(JW 스타일 또는 Normal 스타일)의 DOCX를 사용하세요."
            )

    # 1. 단락별 분류
    lines = []
    title_found = False
    for para in doc.paragraphs:
        text = _docx_normalize_text(para.text)
        if not text:
            continue
        indent = para.paragraph_format.left_indent
        indent_pt = round(indent.pt, 1) if indent else 0.0
        if indent_pt == 0 and para.style and para.style.paragraph_format:
            style_indent = para.style.paragraph_format.left_indent
            if style_indent:
                indent_pt = round(style_indent.pt, 1)
        bold = _docx_first_run_bold(para)
        italic = _docx_any_italic(para)
        style_name = para.style.name if para.style else ""

        tag = _docx_classify_line(text, indent_pt, bold, italic, style_name)

        # 제목은 첫 번째만 인정 (두 번째 볼드 줄은 L1으로 강등)
        if tag == "제목":
            if title_found:
                tag = f"L{_docx_get_level(indent_pt)}"
            else:
                title_found = True

        lines.append({
            "indent_pt": indent_pt, "tag": tag,
            "bold": bold, "italic": italic, "text": text,
        })

    # 2. 후처리
    lines = _docx_merge_split_notes(lines)
    lines = _docx_merge_scripture_linebreaks(lines)

    # 3. raw_lines 포맷 (디버그용)
    raw_lines = []
    for ln in lines:
        flags = []
        if ln["bold"]: flags.append("B")
        if ln["italic"]: flags.append("I")
        flag_str = f" [{','.join(flags)}]" if flags else ""
        raw_lines.append(f"[{ln['indent_pt']:5.1f}pt] [{ln['tag']}]{flag_str} {ln['text']}")

    # 4. 구조화
    result = {
        "title": "",
        "outline_type": None,
        "outline_num": None,
        "version": None,
        "note": None,
        "subtopics": [],
        "raw_lines": raw_lines,
    }

    current_sub = None

    for ln in lines:
        tag = ln["tag"]
        text = ln["text"]

        if tag in ("하단", "제외"):
            # 하단 영역에서 마지막 줄이 버전(예: "1/20")일 수 있음
            vm = re.search(r'No\.?\s*(\d+)\s+(\d+/\d+)', text)
            if vm and not result["version"]:
                result["version"] = vm.group(2)
            continue

        if tag == "제목":
            result["title"] = text
            continue

        if tag == "주의":
            clean = re.sub(r'^(연사의\s*)?유의\s*사항\s*:?\s*', '', text).strip()
            result["note"] = clean or text
            continue

        if tag == "소주제":
            time_minutes = None
            sub_title = text
            dur_m = re.search(r'\(\s*(\d+)\s*분\s*\)', text)
            if dur_m:
                try:
                    time_minutes = int(dur_m.group(1))
                except ValueError:
                    time_minutes = None
                sub_title = re.sub(r'\s*\(\s*\d+\s*분\s*\)\s*', '', text).strip()
            current_sub = {
                "title": sub_title,
                "time_minutes": time_minutes,
                "points": [],
            }
            result["subtopics"].append(current_sub)
            continue

        if tag.startswith("L"):
            try:
                level = int(tag[1:])
            except ValueError:
                continue
            content, scriptures, publications = _docx_extract_refs(text)
            if current_sub is None:
                # 소주제가 없는 10분짜리 골자 등 — 제목을 소주제로 사용
                current_sub = {
                    "title": result["title"] or "",
                    "time_minutes": None,
                    "points": [],
                }
                result["subtopics"].append(current_sub)
            current_sub["points"].append({
                "level": level,
                "number": "",
                "content": content,
                "scriptures": scriptures,
                "publications": publications,
            })

    # 5. 번호 부여
    for sub in result["subtopics"]:
        _docx_assign_numbers(sub["points"])

    return result


# ─── DOCX → 들여쓰기 텍스트 변환 ──────────────────────

_RAW_LINE_RE = re.compile(r'\[\s*[\d.]+pt\]\s*\[([^\]]+)\](?:\s*\[[^\]]*\])?\s*(.*)')


def _lines_to_indented_text(raw_lines: list) -> str:
    """
    parse_outline_docx()의 raw_lines를 [골자 입력] 파서와 호환되는
    들여쓰기 plain text로 변환.

    규칙:
      - [제목] / [유의] / [하단] / [제외] 태그 줄은 제외 (meta로만 사용)
      - [소주제] → 들여쓰기 0칸 ("제목 (N분)" 형태 원문 그대로)
      - [L1] → 들여쓰기 0칸
      - [L2] → 들여쓰기 1칸
      - [L3] → 들여쓰기 2칸
      - [L4] → 들여쓰기 3칸
      - [L5] → 들여쓰기 4칸
      - 본문 텍스트는 원문 그대로 (성구/출판물 괄호 포함, 수정 금지)
    """
    out = []
    for line in raw_lines or []:
        m = _RAW_LINE_RE.match(line)
        if not m:
            continue
        tag = m.group(1).strip()
        text = m.group(2).strip()
        if tag in ('제목', '유의', '하단', '제외'):
            continue
        if tag == '소주제':
            out.append(text)
            continue
        if tag.startswith('L'):
            try:
                level = int(tag[1:])
            except ValueError:
                continue
            if level < 1:
                level = 1
            if level > 5:
                level = 5
            out.append((' ' * (level - 1)) + text)
    return '\n'.join(out)


_VERSION_FALLBACK_RE = re.compile(r'(?<!\d)(\d{1,2}/\d{2})(?!\d)')
_TOTAL_TIME_RE = re.compile(r'(\d+)\s*분\s*(?:을|에)?\s*다룰\s*것')


def _extract_meta_from_docx(parsed: dict, filename: str) -> dict:
    """
    parse_outline_docx() 결과 + 파일명 → meta dict.
    본문 메타 우선, 파일명은 폴백.

    total_time:
      1) 본문 [하단] 줄의 "Nㅎ분에/을 다룰 것" 패턴
      2) 폴백: 각 소주제 time_minutes 합산
    version:
      1) parse_outline_docx 본체가 추출한 값
      2) 파일명 추출
      3) 폴백: [하단] 줄 전수에서 M/YY 패턴 (구버전 포맷 호환)
    """
    fn_meta = parse_outline_filename(filename) if filename else {
        "outline_type": None, "outline_num": None, "outline_year": None, "version": None,
    }
    ot = parsed.get("outline_type") or fn_meta.get("outline_type") or ""
    on = parsed.get("outline_num") or fn_meta.get("outline_num") or ""
    oy = parsed.get("outline_year") or fn_meta.get("outline_year") or ""
    version = parsed.get("version") or fn_meta.get("version") or ""
    title = parsed.get("title") or ""
    note = parsed.get("note") or ""

    total_time = None
    raw_lines = parsed.get("raw_lines", []) or []
    for line in raw_lines:
        m = _RAW_LINE_RE.match(line)
        if not m:
            continue
        tag = m.group(1).strip()
        text = m.group(2).strip()
        if tag == '하단':
            tm = _TOTAL_TIME_RE.search(text)
            if tm:
                try:
                    total_time = int(tm.group(1))
                except ValueError:
                    total_time = None
                break

    if total_time is None:
        times = [s.get("time_minutes") for s in parsed.get("subtopics", []) or []]
        times = [t for t in times if isinstance(t, int) and t > 0]
        if times:
            total_time = sum(times)

    if not version:
        for line in raw_lines:
            m = _RAW_LINE_RE.match(line)
            if not m:
                continue
            tag = m.group(1).strip()
            text = m.group(2).strip()
            if tag != '하단':
                continue
            vm = _VERSION_FALLBACK_RE.search(text)
            if vm:
                version = vm.group(1)
                break

    return {
        "outline_type": ot or None,
        "outline_type_name": (_TYPE_NAMES.get(ot) if ot else None),
        "outline_num": on or None,
        "outline_year": oy or None,
        "version": version or None,
        "title": title or None,
        "note": note or None,
        "total_time": total_time,
    }


# ─── DOCX 파일명 파서 ──────────────────────────────────

def parse_outline_filename(filename: str) -> dict:
    """
    골자 DOCX 파일명에서 outline_type, outline_num, outline_year, version 추출.
    매칭 실패 시 값은 None.

    파일명 규칙:
      S-34_KO_001.docx              → type=S-34,  num=001, year=None
      S-34_KO_001_v09-15.docx       → + version=09/15
      S-31_KO.docx                  → type=S-31,  num=001, year=None
      S-123-26_KO.docx              → type=S-123, num=001, year=26
      S-123-26_KO_v01-26.docx       → + version=01/26
      S-211-26_KO.docx              → type=S-211, num=001, year=26
      CO-26-C_KO.docx               → type=CO_C,  num=001, year=26
      CO-26-C_002_KO.docx           → type=CO_C,  num=002, year=26
      CO-26-R_KO.docx               → type=CO_R,  num=001, year=26
      JWBC-SP_KO_123.docx           → type=JWBC-SP, num=123
    """
    result = {"outline_type": None, "outline_num": None, "outline_year": None, "version": None}
    if not filename:
        return result
    base = os.path.splitext(os.path.basename(filename))[0]
    parts = base.split("_")
    if not parts:
        return result

    # 버전 "_v09-15" 추출
    kept = []
    for p in parts:
        vm = re.match(r'^v(\d+(?:[\-]\d+)*)$', p)
        if vm:
            result["version"] = vm.group(1).replace("-", "/")
        else:
            kept.append(p)

    # 언어 코드 "KO" 제거
    kept = [p for p in kept if p.upper() != "KO"]

    if not kept:
        return result

    first = kept[0]

    # S-XXX 또는 S-XXX-YY (2자리 년도) 형태
    m = re.match(r'^(S-\d+)(?:-(\d{2}))?$', first)
    if m:
        result["outline_type"] = m.group(1)
        if m.group(2):
            result["outline_year"] = m.group(2)
        # num: kept[1:] 에서 첫 숫자 또는 기본값 "001"
        num = None
        for p in kept[1:]:
            if p.isdigit():
                num = p.zfill(3)
                break
        result["outline_num"] = num or "001"
    elif re.match(r'^CO-\d{2}-[CR]$', first):
        # CO-YY-C (순회) 또는 CO-YY-R (지역)
        cm = re.match(r'^CO-(\d{2})-([CR])$', first)
        result["outline_type"] = f"CO_{cm.group(2)}"
        result["outline_year"] = cm.group(1)
        num = None
        for p in kept[1:]:
            if p.isdigit():
                num = p.zfill(3)
                break
        result["outline_num"] = num or "001"
    elif first.startswith("JWBC"):
        result["outline_type"] = first
        for p in kept[1:]:
            if p.isdigit():
                result["outline_num"] = p
                break

    return result
